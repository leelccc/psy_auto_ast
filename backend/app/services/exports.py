from io import BytesIO

from docx import Document
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

# 思源/宋体 CID 字体，保证中文在 PDF 中正常显示。
_CJK_FONT = "STSong-Light"
_INK = HexColor("#1f2329")
_SUBTLE = HexColor("#646a73")
_ACCENT = HexColor("#8a5a4a")


def _register_cjk_font() -> None:
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(_CJK_FONT))
    except Exception:
        # 已注册或环境缺失时忽略，避免影响生成主流程。
        pass


def _escape(text: str) -> str:
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _paragraph_to_html(text: str) -> str:
    """将纯文本转为可在 Paragraph 中安全渲染的 HTML，保留换行。"""
    escaped = _escape(text)
    return "<br/>".join(line for line in escaped.splitlines())


def render_pdf(title: str, content: dict) -> bytes:
    _register_cjk_font()
    output = BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=str(title),
        author="心理档案助手",
    )
    title_style = ParagraphStyle(
        "ReportTitle",
        fontName=_CJK_FONT,
        fontSize=18,
        leading=26,
        spaceAfter=4,
        textColor=_INK,
        alignment=TA_LEFT,
    )
    meta_style = ParagraphStyle(
        "ReportMeta",
        fontName=_CJK_FONT,
        fontSize=9,
        leading=14,
        spaceAfter=10,
        textColor=_SUBTLE,
    )
    heading_style = ParagraphStyle(
        "SectionHeading",
        fontName=_CJK_FONT,
        fontSize=13,
        leading=20,
        spaceBefore=12,
        spaceAfter=6,
        textColor=_ACCENT,
    )
    body_style = ParagraphStyle(
        "SectionBody",
        fontName=_CJK_FONT,
        fontSize=10.5,
        leading=18,
        spaceAfter=4,
        textColor=_INK,
        alignment=TA_LEFT,
    )
    story: list = [
        Paragraph(_escape(title), title_style),
        Paragraph("本文件由心理档案助手生成，供咨询师校订与留痕使用。", meta_style),
    ]
    blocks = content.get("blocks", []) if isinstance(content, dict) else []
    if not blocks:
        story.append(Paragraph("（暂无内容）", body_style))
    for index, block in enumerate(blocks):
        if not isinstance(block, dict):
            continue
        block_title = str(block.get("title", "")).strip()
        block_content = str(block.get("content", "")).strip()
        if not block_title and not block_content:
            continue
        if index > 0:
            story.append(Spacer(1, 4))
        if block_title:
            story.append(Paragraph(_escape(block_title), heading_style))
        if block_content:
            story.append(Paragraph(_paragraph_to_html(block_content), body_style))
    doc.build(story)
    return output.getvalue()


def render_docx(title: str, content: dict) -> bytes:
    document = Document()
    document.add_heading(title, level=0)
    blocks = content.get("blocks", []) if isinstance(content, dict) else []
    if not blocks:
        document.add_paragraph("（暂无内容）")
    for block in blocks:
        if not isinstance(block, dict):
            continue
        block_title = str(block.get("title", "")).strip()
        block_content = str(block.get("content", "")).strip()
        if not block_title and not block_content:
            continue
        if block_title:
            document.add_heading(block_title, level=1)
        if block_content:
            document.add_paragraph(block_content)
    output = BytesIO()
    document.save(output)
    return output.getvalue()
